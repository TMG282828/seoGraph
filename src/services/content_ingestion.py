"""
Content ingestion service for the SEO Content Knowledge Graph System.

This module provides content parsing and ingestion from various file formats
with metadata extraction and preprocessing capabilities.
"""

import asyncio
import hashlib
import mimetypes
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
from urllib.parse import urlparse

import aiofiles
import structlog
from bs4 import BeautifulSoup
import markdown
import pandas as pd

# Optional imports for specific file types
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from config.settings import get_settings
from models.content_models import ContentItem, ContentType, ContentMetadata, ContentMetrics

logger = structlog.get_logger(__name__)


class ContentIngestionError(Exception):
    """Raised when content ingestion fails."""
    pass


class FileProcessor:
    """Base class for file processors."""
    
    def __init__(self):
        """Initialize file processor."""
        self.supported_extensions: List[str] = []
        self.supported_mime_types: List[str] = []
    
    def can_process(self, file_path: str, mime_type: Optional[str] = None) -> bool:
        """
        Check if processor can handle the file.
        
        Args:
            file_path: Path to the file
            mime_type: MIME type of the file
            
        Returns:
            True if processor can handle the file
        """
        extension = Path(file_path).suffix.lower()
        
        if extension in self.supported_extensions:
            return True
        
        if mime_type and mime_type in self.supported_mime_types:
            return True
        
        return False
    
    async def process(self, file_path: str) -> Dict[str, Any]:
        """
        Process file and extract content.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        raise NotImplementedError("Subclasses must implement process method")


class TextFileProcessor(FileProcessor):
    """Processor for plain text files."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.txt', '.md', '.rst']
        self.supported_mime_types = ['text/plain', 'text/markdown']
    
    async def process(self, file_path: str) -> Dict[str, Any]:
        """Process text file."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # Extract title from first line if it looks like a title
            lines = content.split('\n')
            title = None
            
            if lines:
                first_line = lines[0].strip()
                # Check if first line looks like a title (short, no punctuation at end)
                if len(first_line) < 100 and not first_line.endswith('.'):
                    title = first_line
                    content = '\n'.join(lines[1:]).strip()
            
            return {
                'content': content,
                'title': title or Path(file_path).stem,
                'content_type': ContentType.ARTICLE,
                'metadata': {
                    'source_file': file_path,
                    'file_type': 'text',
                    'encoding': 'utf-8'
                }
            }
            
        except Exception as e:
            raise ContentIngestionError(f"Failed to process text file {file_path}: {e}")


class MarkdownFileProcessor(FileProcessor):
    """Processor for Markdown files."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.md', '.markdown']
        self.supported_mime_types = ['text/markdown', 'text/x-markdown']
    
    async def process(self, file_path: str) -> Dict[str, Any]:
        """Process Markdown file."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                markdown_content = await f.read()
            
            # Parse frontmatter if present
            frontmatter = {}
            content = markdown_content
            
            if markdown_content.startswith('---'):
                parts = markdown_content.split('---', 2)
                if len(parts) >= 3:
                    frontmatter_text = parts[1].strip()
                    content = parts[2].strip()
                    
                    # Simple YAML-like parsing for frontmatter
                    for line in frontmatter_text.split('\n'):
                        if ':' in line:
                            key, value = line.split(':', 1)
                            frontmatter[key.strip()] = value.strip().strip('"\'')
            
            # Convert Markdown to HTML for content extraction
            html = markdown.markdown(content)
            soup = BeautifulSoup(html, 'html.parser')
            
            # Extract title
            title = frontmatter.get('title')
            if not title:
                # Look for first heading
                heading = soup.find(['h1', 'h2', 'h3'])
                if heading:
                    title = heading.get_text().strip()
                else:
                    title = Path(file_path).stem
            
            # Extract plain text content
            plain_text = soup.get_text()
            
            # Extract metadata
            metadata = {
                'source_file': file_path,
                'file_type': 'markdown',
                'frontmatter': frontmatter,
                'html_content': html
            }
            
            # Add frontmatter fields to metadata
            if 'description' in frontmatter:
                metadata['meta_description'] = frontmatter['description']
            if 'keywords' in frontmatter:
                metadata['meta_keywords'] = frontmatter['keywords'].split(',')
            if 'tags' in frontmatter:
                metadata['tags'] = frontmatter['tags'].split(',')
            
            return {
                'content': plain_text,
                'title': title,
                'content_type': ContentType.ARTICLE,
                'metadata': metadata
            }
            
        except Exception as e:
            raise ContentIngestionError(f"Failed to process Markdown file {file_path}: {e}")


class HTMLFileProcessor(FileProcessor):
    """Processor for HTML files."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.html', '.htm']
        self.supported_mime_types = ['text/html']
    
    async def process(self, file_path: str) -> Dict[str, Any]:
        """Process HTML file."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                html_content = await f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract title
            title = None
            title_tag = soup.find('title')
            if title_tag:
                title = title_tag.get_text().strip()
            else:
                # Look for h1 tag
                h1_tag = soup.find('h1')
                if h1_tag:
                    title = h1_tag.get_text().strip()
            
            if not title:
                title = Path(file_path).stem
            
            # Extract meta description
            meta_description = None
            meta_desc_tag = soup.find('meta', attrs={'name': 'description'})
            if meta_desc_tag:
                meta_description = meta_desc_tag.get('content')
            
            # Extract meta keywords
            meta_keywords = []
            meta_keywords_tag = soup.find('meta', attrs={'name': 'keywords'})
            if meta_keywords_tag:
                keywords_content = meta_keywords_tag.get('content', '')
                meta_keywords = [kw.strip() for kw in keywords_content.split(',')]
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text content
            text_content = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text_content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text_content = ' '.join(chunk for chunk in chunks if chunk)
            
            metadata = {
                'source_file': file_path,
                'file_type': 'html',
                'meta_description': meta_description,
                'meta_keywords': meta_keywords,
                'html_content': html_content
            }
            
            return {
                'content': text_content,
                'title': title,
                'content_type': ContentType.LANDING_PAGE,
                'metadata': metadata
            }
            
        except Exception as e:
            raise ContentIngestionError(f"Failed to process HTML file {file_path}: {e}")


class PDFFileProcessor(FileProcessor):
    """Processor for PDF files."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.pdf']
        self.supported_mime_types = ['application/pdf']
    
    async def process(self, file_path: str) -> Dict[str, Any]:
        """Process PDF file."""
        if not PYPDF_AVAILABLE:
            raise ContentIngestionError("pypdf library not available for PDF processing")
        
        try:
            # Read PDF content
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                # Extract metadata
                pdf_metadata = pdf_reader.metadata or {}
                title = pdf_metadata.get('/Title', Path(file_path).stem)
                author = pdf_metadata.get('/Author', '')
                subject = pdf_metadata.get('/Subject', '')
                
                # Extract text from all pages
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                # Clean up text
                text_content = re.sub(r'\s+', ' ', text_content).strip()
                
                metadata = {
                    'source_file': file_path,
                    'file_type': 'pdf',
                    'author': author,
                    'subject': subject,
                    'pages': len(pdf_reader.pages),
                    'pdf_metadata': dict(pdf_metadata)
                }
                
                return {
                    'content': text_content,
                    'title': title,
                    'content_type': ContentType.WHITE_PAPER,
                    'metadata': metadata
                }
            
        except Exception as e:
            raise ContentIngestionError(f"Failed to process PDF file {file_path}: {e}")


class DocxFileProcessor(FileProcessor):
    """Processor for DOCX files."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
        self.supported_mime_types = ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']
    
    async def process(self, file_path: str) -> Dict[str, Any]:
        """Process DOCX file."""
        if not DOCX_AVAILABLE:
            raise ContentIngestionError("python-docx library not available for DOCX processing")
        
        try:
            doc = Document(file_path)
            
            # Extract title from document properties or first paragraph
            title = None
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                title = doc.core_properties.title
            elif doc.paragraphs:
                first_para = doc.paragraphs[0].text.strip()
                if len(first_para) < 100:
                    title = first_para
            
            if not title:
                title = Path(file_path).stem
            
            # Extract text content
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Clean up text
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            # Extract metadata
            props = doc.core_properties
            metadata = {
                'source_file': file_path,
                'file_type': 'docx',
                'author': props.author or '',
                'subject': props.subject or '',
                'keywords': props.keywords or '',
                'created': props.created.isoformat() if props.created else None,
                'modified': props.modified.isoformat() if props.modified else None,
                'paragraphs': len(doc.paragraphs)
            }
            
            return {
                'content': text_content,
                'title': title,
                'content_type': ContentType.ARTICLE,
                'metadata': metadata
            }
            
        except Exception as e:
            raise ContentIngestionError(f"Failed to process DOCX file {file_path}: {e}")


class CSVFileProcessor(FileProcessor):
    """Processor for CSV files containing content data."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.csv']
        self.supported_mime_types = ['text/csv', 'application/csv']
    
    async def process(self, file_path: str) -> Dict[str, Any]:
        """Process CSV file."""
        try:
            # Read CSV file
            df = pd.read_csv(file_path)
            
            # Identify content columns (common column names)
            content_columns = ['content', 'text', 'body', 'description', 'article']
            title_columns = ['title', 'name', 'subject', 'headline']
            
            content_col = None
            title_col = None
            
            # Find content and title columns
            for col in df.columns:
                col_lower = col.lower()
                if not content_col and any(cc in col_lower for cc in content_columns):
                    content_col = col
                if not title_col and any(tc in col_lower for tc in title_columns):
                    title_col = col
            
            if not content_col:
                # Use first text column as content
                for col in df.columns:
                    if df[col].dtype == 'object':  # String column
                        content_col = col
                        break
            
            if not content_col:
                raise ContentIngestionError("No suitable content column found in CSV")
            
            # Process each row as separate content item
            content_items = []
            for idx, row in df.iterrows():
                content = str(row[content_col]) if pd.notna(row[content_col]) else ""
                title = str(row[title_col]) if title_col and pd.notna(row[title_col]) else f"Item {idx + 1}"
                
                if content.strip():
                    metadata = {
                        'source_file': file_path,
                        'file_type': 'csv',
                        'row_index': idx,
                        'csv_data': row.to_dict()
                    }
                    
                    content_items.append({
                        'content': content,
                        'title': title,
                        'content_type': ContentType.ARTICLE,
                        'metadata': metadata
                    })
            
            # Return as batch result
            return {
                'batch_items': content_items,
                'title': f"CSV Import: {Path(file_path).stem}",
                'content_type': ContentType.OTHER,
                'metadata': {
                    'source_file': file_path,
                    'file_type': 'csv',
                    'total_rows': len(df),
                    'processed_rows': len(content_items),
                    'columns': list(df.columns)
                }
            }
            
        except Exception as e:
            raise ContentIngestionError(f"Failed to process CSV file {file_path}: {e}")


class ContentIngestionService:
    """
    Service for ingesting content from various file formats.
    
    Provides parsing, metadata extraction, and preprocessing
    for multiple file types.
    """
    
    def __init__(
        self,
        max_file_size_mb: float = 50.0,
        allowed_extensions: Optional[List[str]] = None,
        temp_dir: Optional[str] = None
    ):
        """
        Initialize content ingestion service.
        
        Args:
            max_file_size_mb: Maximum file size in MB
            allowed_extensions: List of allowed file extensions
            temp_dir: Temporary directory for processing
        """
        settings = get_settings()
        
        self.max_file_size_bytes = int(max_file_size_mb * 1024 * 1024)
        self.allowed_extensions = allowed_extensions or [
            '.txt', '.md', '.html', '.htm', '.pdf', '.docx', '.csv'
        ]
        self.temp_dir = temp_dir or '/tmp'
        
        # Initialize file processors
        self.processors = [
            TextFileProcessor(),
            MarkdownFileProcessor(),
            HTMLFileProcessor(),
            PDFFileProcessor(),
            DocxFileProcessor(),
            CSVFileProcessor(),
        ]
        
        logger.info(
            "Content ingestion service initialized",
            max_file_size_mb=max_file_size_mb,
            allowed_extensions=self.allowed_extensions,
            processors_count=len(self.processors)
        )
    
    def _get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get file information and metadata."""
        path = Path(file_path)
        
        # Get file stats
        stat = path.stat()
        
        # Get MIME type
        mime_type, _ = mimetypes.guess_type(file_path)
        
        return {
            'file_path': str(path.absolute()),
            'file_name': path.name,
            'file_stem': path.stem,
            'extension': path.suffix.lower(),
            'size_bytes': stat.st_size,
            'size_mb': stat.st_size / (1024 * 1024),
            'mime_type': mime_type,
            'created_at': datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc),
            'modified_at': datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc),
        }
    
    def _validate_file(self, file_info: Dict[str, Any]) -> None:
        """Validate file for processing."""
        # Check file size
        if file_info['size_bytes'] > self.max_file_size_bytes:
            raise ContentIngestionError(
                f"File too large: {file_info['size_mb']:.1f}MB "
                f"(max: {self.max_file_size_bytes / (1024 * 1024):.1f}MB)"
            )
        
        # Check file extension
        if file_info['extension'] not in self.allowed_extensions:
            raise ContentIngestionError(
                f"File extension '{file_info['extension']}' not allowed. "
                f"Allowed: {', '.join(self.allowed_extensions)}"
            )
    
    def _find_processor(self, file_path: str, mime_type: Optional[str] = None) -> Optional[FileProcessor]:
        """Find appropriate processor for file."""
        for processor in self.processors:
            if processor.can_process(file_path, mime_type):
                return processor
        return None
    
    def _calculate_content_hash(self, content: str) -> str:
        """Calculate hash of content for deduplication."""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()
    
    def _extract_basic_metrics(self, content: str) -> ContentMetrics:
        """Extract basic content metrics."""
        metrics = ContentMetrics()
        
        # Word and character counts
        words = content.split()
        metrics.word_count = len(words)
        metrics.character_count = len(content)
        
        # Paragraph count (simple line-based)
        paragraphs = [p for p in content.split('\n\n') if p.strip()]
        metrics.paragraph_count = len(paragraphs)
        
        # Sentence count (basic period counting)
        sentences = [s for s in content.split('.') if s.strip()]
        metrics.sentence_count = len(sentences)
        
        # Reading time (average 200 words per minute)
        if metrics.word_count > 0:
            metrics.reading_time_minutes = metrics.word_count / 200
        
        return metrics
    
    async def ingest_file(
        self,
        file_path: str,
        tenant_id: str,
        author_id: str,
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> Union[ContentItem, List[ContentItem]]:
        """
        Ingest content from a file.
        
        Args:
            file_path: Path to the file
            tenant_id: Tenant identifier
            author_id: Author user ID
            additional_metadata: Additional metadata to include
            
        Returns:
            ContentItem or list of ContentItems (for batch files like CSV)
        """
        # Get file information
        file_info = self._get_file_info(file_path)
        
        # Validate file
        self._validate_file(file_info)
        
        # Find processor
        processor = self._find_processor(file_path, file_info['mime_type'])
        if not processor:
            raise ContentIngestionError(
                f"No processor available for file type: {file_info['extension']}"
            )
        
        logger.info(
            "Starting file ingestion",
            file_path=file_path,
            file_size_mb=file_info['size_mb'],
            processor=type(processor).__name__
        )
        
        try:
            # Process file
            result = await processor.process(file_path)
            
            # Handle batch results (e.g., CSV with multiple items)
            if 'batch_items' in result:
                content_items = []
                for item_data in result['batch_items']:
                    content_item = await self._create_content_item(
                        item_data, tenant_id, author_id, file_info, additional_metadata
                    )
                    content_items.append(content_item)
                return content_items
            
            # Handle single item
            content_item = await self._create_content_item(
                result, tenant_id, author_id, file_info, additional_metadata
            )
            
            logger.info(
                "File ingestion completed",
                file_path=file_path,
                content_id=content_item.id,
                word_count=content_item.metrics.word_count
            )
            
            return content_item
            
        except Exception as e:
            logger.error(
                "File ingestion failed",
                file_path=file_path,
                error=str(e)
            )
            raise
    
    async def _create_content_item(
        self,
        result: Dict[str, Any],
        tenant_id: str,
        author_id: str,
        file_info: Dict[str, Any],
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> ContentItem:
        """Create ContentItem from processing result."""
        # Extract basic fields
        content = result['content']
        title = result['title']
        content_type = result['content_type']
        
        # Calculate metrics
        metrics = self._extract_basic_metrics(content)
        
        # Prepare metadata
        metadata = ContentMetadata()
        
        # Add file metadata
        file_metadata = result.get('metadata', {})
        if 'meta_description' in file_metadata:
            metadata.meta_description = file_metadata['meta_description']
        if 'meta_keywords' in file_metadata:
            metadata.meta_keywords = file_metadata['meta_keywords']
        
        # Create content item
        content_item = ContentItem(
            title=title,
            content=content,
            content_type=content_type,
            tenant_id=tenant_id,
            author_id=author_id,
            metadata=metadata,
            metrics=metrics,
            source_url=f"file://{file_info['file_path']}",
            custom_fields={
                'file_info': file_info,
                'processing_metadata': file_metadata,
                'content_hash': self._calculate_content_hash(content),
                **(additional_metadata or {})
            }
        )
        
        return content_item
    
    async def ingest_directory(
        self,
        directory_path: str,
        tenant_id: str,
        author_id: str,
        recursive: bool = True,
        progress_callback: Optional[callable] = None
    ) -> List[ContentItem]:
        """
        Ingest all supported files from a directory.
        
        Args:
            directory_path: Directory to process
            tenant_id: Tenant identifier
            author_id: Author user ID
            recursive: Whether to process subdirectories
            progress_callback: Optional progress callback
            
        Returns:
            List of created ContentItems
        """
        directory = Path(directory_path)
        if not directory.is_dir():
            raise ContentIngestionError(f"Directory not found: {directory_path}")
        
        # Find all supported files
        files_to_process = []
        
        if recursive:
            for ext in self.allowed_extensions:
                files_to_process.extend(directory.rglob(f"*{ext}"))
        else:
            for ext in self.allowed_extensions:
                files_to_process.extend(directory.glob(f"*{ext}"))
        
        logger.info(
            "Starting directory ingestion",
            directory=directory_path,
            files_found=len(files_to_process),
            recursive=recursive
        )
        
        content_items = []
        errors = []
        
        for i, file_path in enumerate(files_to_process):
            try:
                result = await self.ingest_file(
                    str(file_path), tenant_id, author_id
                )
                
                # Handle both single items and lists
                if isinstance(result, list):
                    content_items.extend(result)
                else:
                    content_items.append(result)
                
                # Update progress
                if progress_callback:
                    progress = (i + 1) / len(files_to_process)
                    await progress_callback(progress, file_path, len(content_items))
                
            except Exception as e:
                error_info = {
                    'file_path': str(file_path),
                    'error': str(e)
                }
                errors.append(error_info)
                logger.warning(
                    "File ingestion failed in batch",
                    file_path=str(file_path),
                    error=str(e)
                )
        
        logger.info(
            "Directory ingestion completed",
            directory=directory_path,
            content_items_created=len(content_items),
            errors=len(errors)
        )
        
        return content_items
    
    async def ingest_url_content(
        self,
        url: str,
        tenant_id: str,
        author_id: str,
        headers: Optional[Dict[str, str]] = None
    ) -> ContentItem:
        """
        Ingest content from a URL.
        
        Args:
            url: URL to fetch content from
            tenant_id: Tenant identifier
            author_id: Author user ID
            headers: Optional HTTP headers
            
        Returns:
            ContentItem with extracted content
        """
        import httpx
        
        headers = headers or {
            'User-Agent': 'SEO-Content-KnowledgeGraph/1.0'
        }
        
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(url, headers=headers, timeout=30)
                response.raise_for_status()
                
                # Process HTML content
                processor = HTMLFileProcessor()
                
                # Save content to temporary file
                temp_file = Path(self.temp_dir) / f"url_content_{int(time.time())}.html"
                
                async with aiofiles.open(temp_file, 'w', encoding='utf-8') as f:
                    await f.write(response.text)
                
                try:
                    result = await processor.process(str(temp_file))
                    
                    # Add URL metadata
                    result['metadata']['source_url'] = url
                    result['metadata']['http_status'] = response.status_code
                    result['metadata']['content_type'] = response.headers.get('content-type')
                    
                    # Create content item
                    file_info = {
                        'file_path': url,
                        'file_name': urlparse(url).path.split('/')[-1] or 'webpage',
                        'size_bytes': len(response.content),
                        'mime_type': response.headers.get('content-type'),
                        'created_at': datetime.now(timezone.utc),
                        'modified_at': datetime.now(timezone.utc),
                    }
                    
                    content_item = await self._create_content_item(
                        result, tenant_id, author_id, file_info
                    )
                    
                    logger.info(
                        "URL content ingestion completed",
                        url=url,
                        content_id=content_item.id,
                        word_count=content_item.metrics.word_count
                    )
                    
                    return content_item
                    
                finally:
                    # Clean up temporary file
                    if temp_file.exists():
                        temp_file.unlink()
                    
        except Exception as e:
            logger.error("URL content ingestion failed", url=url, error=str(e))
            raise ContentIngestionError(f"Failed to ingest content from URL {url}: {e}")
    
    def get_service_stats(self) -> Dict[str, Any]:
        """Get service statistics."""
        return {
            'max_file_size_mb': self.max_file_size_bytes / (1024 * 1024),
            'allowed_extensions': self.allowed_extensions,
            'processors_available': [type(p).__name__ for p in self.processors],
            'temp_dir': self.temp_dir,
            'pypdf_available': PYPDF_AVAILABLE,
            'docx_available': DOCX_AVAILABLE,
        }


# =============================================================================
# Utility Functions
# =============================================================================

async def ingest_file_simple(
    file_path: str,
    tenant_id: str = "default",
    author_id: str = "system"
) -> ContentItem:
    """
    Simple file ingestion function.
    
    Args:
        file_path: Path to file
        tenant_id: Tenant identifier
        author_id: Author identifier
        
    Returns:
        ContentItem with extracted content
    """
    service = ContentIngestionService()
    return await service.ingest_file(file_path, tenant_id, author_id)


if __name__ == "__main__":
    # Example usage and testing
    async def main():
        service = ContentIngestionService()
        
        # Test with a simple text file
        test_content = """# Test Article

This is a test article for content ingestion.

It has multiple paragraphs and should be processed correctly.
"""
        
        # Create temporary test file
        test_file = Path("/tmp/test_article.md")
        with open(test_file, 'w') as f:
            f.write(test_content)
        
        try:
            # Ingest the file
            content_item = await service.ingest_file(
                str(test_file),
                tenant_id="test-tenant",
                author_id="test-author"
            )
            
            print(f"Ingested content: {content_item.title}")
            print(f"Word count: {content_item.metrics.word_count}")
            print(f"Content type: {content_item.content_type}")
            
            # Service stats
            stats = service.get_service_stats()
            print(f"Service stats: {stats}")
            
        finally:
            # Clean up
            if test_file.exists():
                test_file.unlink()

    asyncio.run(main())