"""
File Upload Ingester for SEO Content Knowledge Graph System.

This module provides comprehensive file upload processing including:
- Support for PDFs, Word docs, text files, and more
- Text extraction from various document formats
- File validation and security scanning
- Batch upload processing
- Metadata extraction and preservation
"""

import asyncio
import logging
from typing import Dict, List, Optional, Any, BinaryIO, Union
from datetime import datetime
import hashlib
import mimetypes
from pathlib import Path
import tempfile
import os

from .base_ingester import BaseIngester, RawContent, ContentSource
from ..database.supabase_client import supabase_client

logger = logging.getLogger(__name__)


class FileProcessor:
    """File processing utilities for different document types."""
    
    @staticmethod
    async def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            import io
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                if page_num >= 100:  # Limit to first 100 pages
                    break
                
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num}: {e}")
            
            return "\n".join(text_parts)
            
        except ImportError:
            logger.warning("PyPDF2 not available, PDF text extraction disabled")
            return "[PDF text extraction requires PyPDF2 library]"
        except Exception as e:
            logger.error(f"Failed to extract PDF text: {e}")
            return "[PDF text extraction failed]"
    
    @staticmethod
    async def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from Word DOCX file."""
        try:
            import docx
            import io
            
            doc = docx.Document(io.BytesIO(file_content))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
            
        except ImportError:
            logger.warning("python-docx not available, DOCX text extraction disabled")
            return "[DOCX text extraction requires python-docx library]"
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            return "[DOCX text extraction failed]"
    
    @staticmethod
    async def extract_text_from_doc(file_content: bytes) -> str:
        """Extract text from legacy Word DOC file."""
        try:
            import tempfile
            import subprocess
            import os
            
            # Use antiword or similar tool for legacy DOC files
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            try:
                # Try antiword first
                result = subprocess.run(['antiword', tmp_file_path], 
                                      capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0:
                    return result.stdout
                else:
                    # Fallback to catdoc
                    result = subprocess.run(['catdoc', tmp_file_path], 
                                          capture_output=True, text=True, timeout=30)
                    if result.returncode == 0:
                        return result.stdout
                
                return "[DOC text extraction failed - no suitable converter found]"
                
            finally:
                os.unlink(tmp_file_path)
                
        except Exception as e:
            logger.error(f"Failed to extract DOC text: {e}")
            return "[DOC text extraction failed]"
    
    @staticmethod
    async def extract_text_from_rtf(file_content: bytes) -> str:
        """Extract text from RTF file."""
        try:
            from striprtf.striprtf import rtf_to_text
            
            rtf_content = file_content.decode('utf-8', errors='ignore')
            return rtf_to_text(rtf_content)
            
        except ImportError:
            logger.warning("striprtf not available, RTF text extraction disabled")
            return "[RTF text extraction requires striprtf library]"
        except Exception as e:
            logger.error(f"Failed to extract RTF text: {e}")
            return "[RTF text extraction failed]"
    
    @staticmethod
    async def extract_text_from_txt(file_content: bytes) -> str:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            return file_content.decode('utf-8', errors='replace')
            
        except Exception as e:
            logger.error(f"Failed to extract text from text file: {e}")
            return "[Text file extraction failed]"
    
    @staticmethod
    async def extract_text_from_csv(file_content: bytes) -> str:
        """Extract text from CSV file."""
        try:
            import csv
            import io
            
            # Decode content
            csv_content = file_content.decode('utf-8', errors='replace')
            
            # Parse CSV and convert to readable text
            text_parts = []
            csv_reader = csv.reader(io.StringIO(csv_content))
            
            for row_num, row in enumerate(csv_reader):
                if row_num == 0:
                    # Header row
                    text_parts.append("CSV Headers: " + " | ".join(row))
                elif row_num < 100:  # Limit to first 100 rows
                    # Data rows
                    non_empty_cells = [cell for cell in row if cell.strip()]
                    if non_empty_cells:
                        text_parts.append("Row {}: {}".format(row_num + 1, " | ".join(non_empty_cells)))
                else:
                    break
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Failed to extract CSV text: {e}")
            return "[CSV text extraction failed]"


class FileUploadIngester(BaseIngester):
    """
    File upload ingester with comprehensive document processing.
    
    Features:
    - Support for multiple document formats
    - Text extraction from PDFs, Word docs, etc.
    - File validation and security scanning
    - Batch processing capabilities
    - Metadata extraction and preservation
    - File deduplication
    """
    
    def __init__(self, organization_id: str, upload_path: str = "/tmp/uploads"):
        super().__init__("file_upload", organization_id)
        self.upload_path = Path(upload_path)
        self.upload_path.mkdir(parents=True, exist_ok=True)
        
        # Supported file types and their processors
        self.file_processors = {
            'application/pdf': FileProcessor.extract_text_from_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': FileProcessor.extract_text_from_docx,
            'application/msword': FileProcessor.extract_text_from_doc,
            'application/rtf': FileProcessor.extract_text_from_rtf,
            'text/plain': FileProcessor.extract_text_from_txt,
            'text/csv': FileProcessor.extract_text_from_csv,
            'text/markdown': FileProcessor.extract_text_from_txt,
            'text/html': FileProcessor.extract_text_from_txt,
        }
        
        # File size limits (in bytes)
        self.max_file_sizes = {
            'application/pdf': 50 * 1024 * 1024,  # 50MB
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 25 * 1024 * 1024,  # 25MB
            'application/msword': 25 * 1024 * 1024,  # 25MB
            'text/plain': 10 * 1024 * 1024,  # 10MB
            'default': 20 * 1024 * 1024  # 20MB default
        }
    
    async def validate_source(self, source_config: Dict[str, Any]) -> bool:
        """Validate file upload configuration."""
        try:
            # Check upload directory exists and is writable
            if not self.upload_path.exists():
                self.upload_path.mkdir(parents=True, exist_ok=True)
            
            # Test write permissions
            test_file = self.upload_path / "test_write.tmp"
            test_file.write_text("test")
            test_file.unlink()
            
            return True
            
        except Exception as e:
            self.logger.error(f"File upload validation failed: {e}")
            return False
    
    async def extract_content(self, source_config: Dict[str, Any]) -> List[RawContent]:
        """Extract content from uploaded files."""
        uploaded_files = source_config.get('uploaded_files', [])
        
        if not uploaded_files:
            self.logger.warning("No files provided for extraction")
            return []
        
        raw_contents = []
        
        for file_info in uploaded_files:
            try:
                content = await self._process_uploaded_file(file_info)
                if content:
                    raw_contents.append(content)
            except Exception as e:
                self.logger.error(f"Failed to process uploaded file {file_info.get('filename', 'unknown')}: {e}")
        
        self.logger.info(f"Successfully processed {len(raw_contents)} uploaded files")
        return raw_contents
    
    async def process_file_upload(self, file_data: bytes, filename: str, 
                                mime_type: Optional[str] = None) -> Dict[str, Any]:
        """Process a single file upload."""
        try:
            # Detect MIME type if not provided
            if not mime_type:
                mime_type, _ = mimetypes.guess_type(filename)
                if not mime_type:
                    mime_type = 'application/octet-stream'
            
            # Validate file
            validation_result = await self._validate_file(file_data, filename, mime_type)
            if not validation_result['valid']:
                return validation_result
            
            # Save file temporarily
            temp_file_path = await self._save_temp_file(file_data, filename)
            
            try:
                # Process file
                file_info = {
                    'filename': filename,
                    'mime_type': mime_type,
                    'file_path': str(temp_file_path),
                    'file_size': len(file_data),
                    'uploaded_at': datetime.now().isoformat()
                }
                
                raw_content = await self._process_uploaded_file(file_info)
                
                if raw_content:
                    return {
                        'success': True,
                        'content_id': raw_content.content_id,
                        'title': raw_content.title,
                        'word_count': raw_content.word_count,
                        'file_size': raw_content.file_size,
                        'content_type': raw_content.content_type
                    }
                else:
                    return {
                        'success': False,
                        'error': 'Failed to extract content from file'
                    }
                
            finally:
                # Clean up temporary file
                if temp_file_path.exists():
                    temp_file_path.unlink()
        
        except Exception as e:
            self.logger.error(f"Failed to process file upload {filename}: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def process_batch_upload(self, files: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Process multiple file uploads in batch."""
        results = {
            'success': True,
            'total_files': len(files),
            'processed_files': 0,
            'failed_files': 0,
            'results': [],
            'errors': []
        }
        
        for file_info in files:
            try:
                file_data = file_info['data']
                filename = file_info['filename']
                mime_type = file_info.get('mime_type')
                
                result = await self.process_file_upload(file_data, filename, mime_type)
                
                if result['success']:
                    results['processed_files'] += 1
                else:
                    results['failed_files'] += 1
                    results['errors'].append({
                        'filename': filename,
                        'error': result['error']
                    })
                
                results['results'].append(result)
                
            except Exception as e:
                results['failed_files'] += 1
                results['errors'].append({
                    'filename': file_info.get('filename', 'unknown'),
                    'error': str(e)
                })
        
        if results['failed_files'] > 0:
            results['success'] = results['processed_files'] > results['failed_files']
        
        return results
    
    async def _validate_file(self, file_data: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Validate uploaded file."""
        # Check file size
        file_size = len(file_data)
        max_size = self.max_file_sizes.get(mime_type, self.max_file_sizes['default'])
        
        if file_size > max_size:
            return {
                'valid': False,
                'error': f'File size {file_size} exceeds maximum allowed size {max_size}'
            }
        
        # Check if file type is supported
        if mime_type not in self.file_processors:
            return {
                'valid': False,
                'error': f'Unsupported file type: {mime_type}'
            }
        
        # Check filename
        if not filename or len(filename) > 255:
            return {
                'valid': False,
                'error': 'Invalid filename'
            }
        
        # Basic security checks
        dangerous_extensions = ['.exe', '.bat', '.cmd', '.sh', '.ps1', '.vbs', '.js']
        file_ext = Path(filename).suffix.lower()
        
        if file_ext in dangerous_extensions:
            return {
                'valid': False,
                'error': f'Potentially dangerous file extension: {file_ext}'
            }
        
        # Check for malicious content patterns
        if await self._scan_for_malicious_content(file_data):
            return {
                'valid': False,
                'error': 'File contains potentially malicious content'
            }
        
        return {'valid': True}
    
    async def _scan_for_malicious_content(self, file_data: bytes) -> bool:
        """Basic malicious content scanning."""
        try:
            # Check for common malicious patterns
            malicious_patterns = [
                b'<script',
                b'javascript:',
                b'vbscript:',
                b'onload=',
                b'onerror=',
                b'eval(',
                b'document.cookie',
                b'window.location'
            ]
            
            file_data_lower = file_data.lower()
            
            for pattern in malicious_patterns:
                if pattern in file_data_lower:
                    return True
            
            return False
            
        except Exception as e:
            self.logger.warning(f"Failed to scan for malicious content: {e}")
            return False
    
    async def _save_temp_file(self, file_data: bytes, filename: str) -> Path:
        """Save file temporarily for processing."""
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_hash = hashlib.md5(file_data).hexdigest()[:8]
        safe_filename = "".join(c for c in filename if c.isalnum() or c in '._-')
        
        temp_filename = f"{timestamp}_{file_hash}_{safe_filename}"
        temp_file_path = self.upload_path / temp_filename
        
        # Save file
        with open(temp_file_path, 'wb') as f:
            f.write(file_data)
        
        return temp_file_path
    
    async def _process_uploaded_file(self, file_info: Dict[str, Any]) -> Optional[RawContent]:
        """Process a single uploaded file."""
        try:
            filename = file_info['filename']
            mime_type = file_info['mime_type']
            file_path = file_info.get('file_path')
            
            # Read file content
            if file_path and Path(file_path).exists():
                with open(file_path, 'rb') as f:
                    file_content = f.read()
            else:
                # If file_path not available, assume content is in file_info
                file_content = file_info.get('content', b'')
            
            # Extract text based on file type
            processor = self.file_processors.get(mime_type)
            if not processor:
                self.logger.warning(f"No processor available for {mime_type}")
                return None
            
            text_content = await processor(file_content)
            
            if not text_content or len(text_content.strip()) < 50:
                self.logger.info(f"Insufficient content extracted from {filename}")
                return None
            
            # Generate content ID
            content_id = self._generate_content_id("file_upload", filename)
            
            # Extract title from filename
            title = Path(filename).stem
            
            # Calculate metadata
            word_count = self._count_words(text_content) if not text_content.startswith('[') else 0
            content_hash = self._calculate_content_hash(text_content)
            
            return RawContent(
                content_id=content_id,
                source_id="file_upload",  # Will be updated with actual source ID
                raw_text=text_content,
                content_type=self._get_content_type_from_mime(mime_type),
                title=title,
                url=None,
                metadata={
                    "original_filename": filename,
                    "mime_type": mime_type,
                    "file_size": len(file_content),
                    "uploaded_at": file_info.get('uploaded_at', datetime.now().isoformat()),
                    "processing_method": processor.__name__,
                    "requires_processing": text_content.startswith('[')
                },
                content_hash=content_hash,
                word_count=word_count,
                file_size=len(file_content)
            )
            
        except Exception as e:
            self.logger.error(f"Failed to process uploaded file {file_info.get('filename', 'unknown')}: {e}")
            return None
    
    def _get_content_type_from_mime(self, mime_type: str) -> str:
        """Get content type from MIME type."""
        mime_to_content_type = {
            'application/pdf': 'pdf_document',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'word_document',
            'application/msword': 'word_document',
            'application/rtf': 'rtf_document',
            'text/plain': 'text_file',
            'text/csv': 'spreadsheet',
            'text/markdown': 'markdown_file',
            'text/html': 'html_file'
        }
        
        return mime_to_content_type.get(mime_type, 'unknown_document')
    
    async def create_file_upload_source(self, source_name: str, source_config: Dict[str, Any]) -> ContentSource:
        """Create a new file upload content source."""
        try:
            source_id = f"file_upload_{hashlib.md5(source_name.encode()).hexdigest()[:16]}"
            
            source = ContentSource(
                source_id=source_id,
                source_type="file_upload",
                source_url=None,
                source_metadata={
                    "source_name": source_name,
                    "config": source_config,
                    "upload_stats": {
                        "total_files_uploaded": 0,
                        "total_files_processed": 0,
                        "total_size_processed": 0
                    }
                },
                organization_id=self.organization_id
            )
            
            # Store in database
            await self._store_content_source(source)
            
            return source
            
        except Exception as e:
            self.logger.error(f"Failed to create file upload source: {e}")
            raise
    
    async def _store_content_source(self, source: ContentSource):
        """Store content source in database."""
        try:
            data = {
                "source_id": source.source_id,
                "source_type": source.source_type,
                "source_url": source.source_url,
                "source_metadata": json.dumps(source.source_metadata),
                "organization_id": source.organization_id,
                "created_at": source.created_at.isoformat(),
                "last_crawled": source.last_crawled.isoformat() if source.last_crawled else None,
                "is_active": source.is_active
            }
            
            supabase_client.client.table("content_sources").upsert(data, on_conflict="source_id").execute()
            
        except Exception as e:
            self.logger.error(f"Failed to store content source: {e}")
            raise
    
    async def get_upload_statistics(self, source_id: str) -> Dict[str, Any]:
        """Get upload statistics for a source."""
        try:
            # Get raw content statistics
            raw_result = supabase_client.client.table("raw_content").select("file_size, word_count, content_type").eq("source_id", source_id).eq("organization_id", self.organization_id).execute()
            
            # Get processed content statistics
            processed_result = supabase_client.client.table("processed_content").select("quality_score, processing_status").eq("source_id", source_id).eq("organization_id", self.organization_id).execute()
            
            raw_data = raw_result.data
            processed_data = processed_result.data
            
            # Calculate statistics
            stats = {
                "total_files": len(raw_data),
                "total_size": sum(item.get("file_size", 0) for item in raw_data),
                "total_words": sum(item.get("word_count", 0) for item in raw_data),
                "processed_files": len([item for item in processed_data if item.get("processing_status") == "completed"]),
                "failed_files": len([item for item in processed_data if item.get("processing_status") == "failed"]),
                "average_quality_score": 0,
                "content_type_distribution": {}
            }
            
            # Calculate average quality score
            quality_scores = [item.get("quality_score", 0) for item in processed_data if item.get("quality_score")]
            if quality_scores:
                stats["average_quality_score"] = sum(quality_scores) / len(quality_scores)
            
            # Calculate content type distribution
            content_types = [item.get("content_type", "unknown") for item in raw_data]
            for content_type in content_types:
                stats["content_type_distribution"][content_type] = stats["content_type_distribution"].get(content_type, 0) + 1
            
            return stats
            
        except Exception as e:
            self.logger.error(f"Failed to get upload statistics: {e}")
            return {}